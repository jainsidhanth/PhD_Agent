"""Convert generated Markdown (CV, SOP, proposal) to .docx using python-docx."""
import io
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BURNT = RGBColor(0xA6, 0x4B, 0x2A)


def _add_runs_with_bold(paragraph, text):
    """Handle **bold** inline segments."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(md_text, title=None):
    """Return docx bytes for the given markdown string."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if title:
        h = doc.add_heading(title, level=0)

    lines = md_text.replace("\r\n", "\n").split("\n")
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue

        if line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            for r in h.runs:
                r.font.color.rgb = BURNT
        elif line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            for r in h.runs:
                r.font.color.rgb = BURNT
        elif line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            for r in h.runs:
                r.font.color.rgb = BURNT
        elif re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, text)
        elif re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_bold(p, text)
        else:
            p = doc.add_paragraph()
            _add_runs_with_bold(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
